import sys
import os
import base64
import argparse
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from lxml import etree
from datetime import date
import zipfile
from pathlib import Path
from io import BytesIO

# TODO: 
# - page breaks

def extract_images_and_map_relationships(docx_path):
    """
    Extract images from a DocX file and map relationship IDs to image filenames.
    """

    # Open the DocX file as a ZIP archive
    with zipfile.ZipFile(docx_path) as docx_zip:
        # Extract all files in the 'word/media' directory
        image_files = [f for f in docx_zip.namelist() if f.startswith("word/media/")]

        # Save images and map their relationship IDs
        for idx, image_file in enumerate(image_files):
            # Get the image filename and extension
            image_filename = os.path.basename(image_file)

            # Save the image to the output directory
            output_path = os.path.join(image_dir, image_filename)

            if not os.path.exists(image_dir):
                os.makedirs(image_dir)

            with docx_zip.open(image_file) as image_data:
                with open(output_path, "wb") as img_file:
                    img_file.write(image_data.read())

            # Map the relationship ID to the saved image filename
            relationship_id = os.path.splitext(image_filename)[0]
            image_map[relationship_id] = image_filename

        # Parse the document relationships to map relationship IDs to image paths
        rels_path = "word/_rels/document.xml.rels"
        if rels_path in docx_zip.namelist():
            with docx_zip.open(rels_path) as rels_file:
                rels_xml = rels_file.read()
                rels_root = etree.fromstring(rels_xml)

                for rel in rels_root.findall(".//{http://schemas.openxmlformats.org/package/2006/relationships}Relationship"):
                    if "image" in rel.get("Type"):
                        relationship_id = rel.get("Id")
                        target = rel.get("Target")
                        if target.startswith("media/"):
                            image_filename = os.path.basename(target)
                            image_map[relationship_id] = image_filename

    return image_map

def get_color(run):
    """Extract font color from the run."""
    if run.font.color and run.font.color.rgb:
        return f"color: #{run.font.color.rgb};"
    if run.style and run.style.font.color and run.style.font.color.rgb:
        return f"color: #{run.style.font.color.rgb};"
    return ""

def debug_elm(elm):
    for attr in dir(elm):
        if not attr.startswith("_") and not callable(getattr(elm, attr)):
            value = getattr(elm, attr)
            print(f"{attr}: {value}")

def get_font_size(run):
    """Extract font size in points."""
    if run.font.size:
        return f"font-size: {run.font.size.pt}pt;"
    if run.style and run.style.font.size:
        return f"font-size: {run.style.font.size.pt}pt;"
    return  ""

def get_text_styles(run):
    """Extract text styles like bold, italic, underline, superscript."""
    styles = []
    if run.bold or ( run.style and run.style.font.bold ):
        styles.append("font-weight: bold;")
    if run.italic or ( run.style and run.style.font.italic ):
        styles.append("font-style: italic;")
    if run.underline or ( run.style and run.style.font.underline ):
        styles.append("text-decoration: underline;")
    if run.font.superscript or ( run.style and run.style.font.superscript ):
        styles.append("vertical-align: super; font-size: smaller;")
    if run.font.subscript or ( run.style and run.style.font.subscript ):
        styles.append("vertical-align: sub; font-size: smaller;")
    return " ".join(styles)

def paragraph_to_css(para):
    """Converts a docx paragraph's style to CSS."""
    css = []

    # Alignment
    align_map = {
        None: "left",
        0: "left",   # WD_ALIGN_PARAGRAPH.LEFT
        1: "center", # WD_ALIGN_PARAGRAPH.CENTER
        2: "right",  # WD_ALIGN_PARAGRAPH.RIGHT
        3: "justify" # WD_ALIGN_PARAGRAPH.JUSTIFY
    }
    alignment = align_map.get(para.paragraph_format.alignment, "left")
    css.append(f"text-align: {alignment};")

    # Indentation (left, right, first-line)
    if para.paragraph_format.left_indent:
        css.append(f"margin-left: {para.paragraph_format.left_indent.pt}pt;")
    if para.paragraph_format.right_indent:
        css.append(f"margin-right: {para.paragraph_format.right_indent.pt}pt;")
    if para.paragraph_format.first_line_indent:
        css.append(f"text-indent: {para.paragraph_format.first_line_indent.pt}pt;")

    # Spacing (before, after, line spacing)
    if para.paragraph_format.space_before:
        css.append(f"margin-top: {para.paragraph_format.space_before.pt}pt;")
    if para.paragraph_format.space_after:
        css.append(f"margin-bottom: {para.paragraph_format.space_after.pt}pt;")
    if para.paragraph_format.line_spacing:
        line_spacing = para.paragraph_format.line_spacing
        line_spacing_rule = para.paragraph_format.line_spacing_rule
        line_height = normalize_line_height(line_spacing, line_spacing_rule)
        css.append(f"line-height: {line_height};")
    
    """Extract background color from a paragraph's XML structure."""
    shd = para._element.find(".//w:shd", namespaces)
    if shd is not None: 
        fill = shd.get(qn("w:fill"))  # Get the w:fill attribute
        if fill and fill != 'FFFFFF':
            # Convert DOCX color (e.g., "FF0000") to CSS format (e.g., "#FF0000")
            css.append(f"background-color: #{fill}")
    pPr = para._element.pPr  # Access paragraph properties

    return " ".join(css)

def normalize_line_height(line_spacing, line_spacing_rule):
    """
    Normalize the line height value from DOCX to a reasonable CSS value.
    """
    if line_spacing_rule == 4:  # WD_LINE_SPACING.EXACTLY
        # Convert from twentieths of a point to points
        return line_spacing / 20
    elif line_spacing_rule == 5:  # WD_LINE_SPACING.MULTIPLE
        # Return the multiplier directly
        return line_spacing
    else:
        # Default line height (e.g., single spacing)
        return 1.2

def get_background_color(para):
    return ""

def get_tag(element):
    tagname = element.tag
    for ns in namespaces:
        nst = namespaces[ns]
        tagname = tagname.replace("{"+nst+"}", "")
    return tagname

def process_run(run):
    """Convert a run (styled text) to TEI <hi> with a single style attribute."""
    text_elem = etree.Element("hi")

    # Collect all styles
    styles = get_color(run) + get_font_size(run) + get_text_styles(run)
    if styles:
        text_elem.set("style", styles.strip())

    xml_str = run._element.xml
    if "w:footnoteReference" in xml_str:
        # Extract footnote ID
        run_xml = etree.fromstring(xml_str)
        for footnote_ref in run_xml.findall(".//w:footnoteReference", namespaces):
            footnote_id = footnote_ref.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id")
            if footnote_id in footnote_map:
                # create a <note> with the text inside for the footnote
                text_elem = etree.Element("note")
                text_elem.set("id", "fn-" + str(footnote_id))
                text_elem.text = footnote_map.get(footnote_id)
    else:
        text_elem.text = run.text

    return text_elem

def append_mixed_content(src, dst):
    """ Append all children and text content from `src` to `dst` in the correct order. """

    # Find the last node in the destination (last child OR text)
    last_child = dst[-1] if len(dst) > 0 else None

    # If there's existing text and no elements, append to the text
    if last_child is None:
        dst.text = (dst.text or "") + (src.text or "")
    else:
        # If there's already an element, append new text to its tail
        last_child.tail = (last_child.tail or "") + (src.text or "")

    # Append each child while preserving order
    for child in src:
        copied_child = copy.deepcopy(child)  # Copy child element
        dst.append(copied_child)  # Append to destination

        if child.tail:
            copied_child.tail = child.tail  # Preserve tail text after each element
            
def process_paragraph(para):
    """Convert a paragraph to TEI."""


    para_elem = etree.Element("p")
    para_elem.tail = "\n"

    para_style = paragraph_to_css(para)
    if para_style:
        para_elem.set("style", para_style)

    # Extract paragraph background color
    background_color = get_background_color(para)
    if background_color:
        para_elem.set("style", background_color.strip())
    
    run_map = {}
    for run in para.runs:
        run_map[run._element] = run
    
    lasthi = para_elem
    for child in para._element:
        tagname = get_tag(child)
        if tagname in [ "pPr", "proofErr", "bookmarkStart", "bookmarkEnd", "smartTag" ] : # things that can be ignored
            pass
        elif tagname == "hyperlink": # Hyperlinks
            ns = '{' + namespaces['r'] + '}'
            rId = child.get(f"{ns}id")
            if rId in hyperlink_map:
                url = hyperlink_map[rId]
                ref_elem = etree.Element("ref")
                ref_elem.set("target", url)
                hyperlinktext = ""
                for run_xml in child.findall(".//w:t", namespaces=namespaces):
                    hyperlinktext = hyperlinktext + run_xml.text
                ref_elem.text = hyperlinktext # This should parse a full run inside a <w:hyperlink> if it exists
                para_elem.append(ref_elem)
        elif tagname == "r": # Styled elements
            r_id = child.getparent().get("r:id")
            hi = process_run(run_map[child])
            if len(hi) > 0 or ( hi.text is not None and hi.text):
                # Skip empty elements
                if not hi.attrib: 
                    # Add to the paragraph directly if there is no styling
                    append_mixed_content(hi, para_elem)
                    lasthi = para_elem
                elif hi.get("style") == lasthi.get("style"):
                    # Add to the last hi if styling has not changed
                    append_mixed_content(hi, lasthi)
                else:
                    # Add the new hi element
                    para_elem.append(hi)
                    lasthi = hi      
        else:
            print("Unhandled paragraph child: ", tagname)          

        # Handle images
        for drawing in child.findall(".//w:drawing", namespaces=namespaces):
            blip = drawing.find(".//a:blip", namespaces=namespaces)
            if blip is not None:
                relationship_id = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                if relationship_id:
                    # Find the corresponding image filename
                    image_filename = image_map.get(relationship_id)
                    if image_filename:
                        # Add a <figure> and <graphic> element to the TEI
                        figure = etree.SubElement(para_elem, "figure")
                        figure.set("id", relationship_id)
                        graphic = etree.SubElement(figure, "graphic", url=os.path.join(image_reldir, image_filename))
    
    # Skip the paragraph if it is empty
    if len(para_elem) == 0 and not para_elem.text:
        return None

    return para_elem

def process_table(table):
    """Convert a DOCX table to TEI."""
    table_elem = etree.Element("table")
    
    for row in table.rows:
        row_elem = etree.Element("row")
        for cell in row.cells:
            cell_elem = etree.Element("cell")
            for para in cell.paragraphs:
                processed_para = process_paragraph(para)
                if processed_para is not None:  # Explicit check
                    cell_elem.append(processed_para)
            row_elem.append(cell_elem)
        table_elem.append(row_elem)

    return table_elem

def extract_hyperlinks(doc):
    for rId in doc.part.rels:
        rel = doc.part.rels[rId]
        if hasattr(rel, "target_ref"):
            hyperlink_map[rId] = rel.target_ref

def extract_footnotes(doc):
    """Extract footnotes manually from docx XML."""
    footnotes_elem = etree.Element("notes")
 
    footnotes_part = None
    for rel in doc.part.rels.values():
        if "footnotes" in rel.target_ref:
            footnotes_part = rel.target_part
            break
            
    if not footnotes_part:
        return

    fncnt = 1
    footnotes_xml = etree.parse(BytesIO(footnotes_part.blob))
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    for footnote in footnotes_xml.findall("w:footnote", ns):
        footnote_id = footnote.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id")
        fntext = "".join(node.text or "" for node in footnote.findall(".//w:t", ns))
        note_elem = etree.Element("note")
        note_elem.text = fntext
        footnote_id = footnote.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id")
        note_elem.set("id", "fn-" + footnote_id)
        fncnt = fncnt + 1
        footnotes_elem.append(note_elem)
        footnote_map[footnote_id] = fntext

    # Skip the paragraph if it is empty
    if len(footnotes_elem) == 0 and not footnotes_elem.text:
        return None

    return footnotes_elem

def convert_docx_to_tei(docx_file, output_tei_file):

    try:
        # Attempt to open the .docx file
        doc = Document(docx_file)
    except PackageNotFoundError:
        print(f"Error: The file '{docx_file}' is not a valid .docx file or is corrupted.")
        sys.exit(1)
    except FileNotFoundError:
        print(f"Error: The file '{docx_file}' does not exist.")
        sys.exit(1)
    except Exception as e:
        print(f"An unexpected error occurred: {e}")
        sys.exit(1)

    image_map = extract_images_and_map_relationships(docx_file)
    footnotes = extract_footnotes(doc)
    hyperlinks = extract_hyperlinks(doc)

    XML_NS = "http://www.w3.org/XML/1998/namespace"

    # Create TEI structure
    tei = etree.Element("TEI")
    tei_header = etree.SubElement(tei, "teiHeader")
    text = etree.SubElement(tei, "text")
    text.set(f"{{{XML_NS}}}space", "preserve")
    text.set("id", os.path.splitext(os.path.basename(docx_file))[0])
    body = etree.SubElement(text, "body")


    filedesc = etree.SubElement(tei_header, "fileDesc")
    notesstmt = etree.SubElement(filedesc, "notesStmt")
    note = etree.SubElement(notesstmt, "note")
    note.set("n", "orgfile")
    note.text = orgfile
    revisiondesc = etree.SubElement(tei_header, "revisionDesc")
    change = etree.SubElement(revisiondesc, "change")
    change.set("who", "docx2tei")
    today = str(date.today())
    change.set("when", today)
    change.text = "Converted from DOCX file " + docx_file
    titlestmt = etree.SubElement(filedesc, "titleStmt")
    profiledesc = etree.SubElement(tei_header, "profileDesc")
    
    # Handle the document metadata
    metadata = doc.core_properties
    if doc.core_properties.title:
        title = etree.SubElement(titlestmt, "title")
        title.text = doc.core_properties.title
    if doc.core_properties.author:
        author = etree.SubElement(titlestmt, "author")
        author.text = doc.core_properties.author
    if doc.core_properties.created:
        cdate = etree.SubElement(titlestmt, "date")
        cdate.text = str(doc.core_properties.created)
    if doc.core_properties.keywords:
        textclass = etree.SubElement(profiledesc, "textClass")
        keywords = etree.SubElement(textclass, "keywords")
        term = etree.SubElement(keywords, "term")
        term.text = doc.core_properties.keywords
    if doc.core_properties.language:
        langusage = etree.SubElement(profiledesc, "langUsage")
        language = etree.SubElement(langusage, "language")
        language.text = doc.core_properties.language


    # Map the docx elements onto a map
    para_map = {}
    for para in doc.paragraphs:
        para_map[para._element] = para  # Map lxml element to python-docx paragraph
    table_map = {}
    for table in doc.tables:
        table_map[table._element] = table  # Map lxml element to python-docx paragraph

    # Iterate through the document body in order
    for element in doc._element.findall(".//w:body/*", namespaces=namespaces):
        if element.tag.endswith("p"):  # Paragraph
            processed_element = process_paragraph(para_map.get(element))
        elif element.tag.endswith("tbl"):  # Table
            processed_element = process_table(table_map.get(element))
        elif element.tag.endswith("sectPr"):  # Table
            # Section Properties - skip
            pass
        else:
            print("Unknown: ", element.tag)
            
        if processed_element is not None:  # Explicit check
            body.append(processed_element)

    # Place footnotes (endnotes) - make it optional!
#     if footnotes is not None:
#         text.append(footnotes)

    # Save TEI XML
    tree = etree.ElementTree(tei)
    tree.write(output_tei_file, pretty_print=True, xml_declaration=True, encoding="utf-8")

    print(f"Conversion complete! Saved as {output_tei_file}")

parser = argparse.ArgumentParser(description="Convert a DOCX file to a (TEITOK style) TEI/XML file")
parser.add_argument("file", help="DOCX file to convert")
parser.add_argument("--force", help="force when output file exists", action="store_true")
parser.add_argument("--debug", help="debug mode", action="store_true")
parser.add_argument("-o", "--output", help="ouput TEI filename", type=str, default="")
parser.add_argument("-i", "--image_dir", help="directory to store the images in", type=str, default="")
parser.add_argument("--orgfile", help="original file that the DOCX was created from", type=str, default="")
args = parser.parse_args()


docx_filename = args.file
fileid = os.path.splitext(os.path.basename(docx_filename))[0]

# Command-line usage
if not os.path.isfile(docx_filename):
    print("No such file: ", docx_filename)
    sys.exit(1)

# Determine default TEITOK style filenames or default to generic names
tei_filename = args.output
if not tei_filename:
    index = docx_filename.find("Originals")
    if index != -1:
        tei_filename = docx_filename[:index] + "xmlfiles/" + fileid + ".xml"
        fileid = os.path.splitext(os.path.basename(docx_filename))[0]
    else:
        tei_filename = os.path.splitext(docx_filename)[0] + ".xml"

xmlid = os.path.splitext(os.path.basename(tei_filename))[0]

image_dir = args.image_dir 
if not image_dir:
    index = tei_filename.find("xmlfiles")
    if index != -1:
        image_dir = tei_filename[:index] + "Graphics/" + xmlid
    else:
        image_dir = os.path.splitext(tei_filename)[0] + "_files"

orgfile = args.orgfile
if not orgfile:
    orgfile = docx_filename

# Determine what to put as the link for images
path = Path(image_dir)
if path.parts[1:] and path.parts[-2] == "Graphics":
    image_reldir = path.parts[-1]
else:
    image_reldir = image_dir
    
# Define global variables to hold the image and footnote references
image_map = {}
footnote_map = {}
hyperlink_map  = {}

namespaces = {
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
}

convert_docx_to_tei(docx_filename, tei_filename)
